from __future__ import annotations

import json
import logging
import os
import uuid
import zipfile
from dataclasses import dataclass
from io import BytesIO
from typing import Dict, List, Optional, Tuple

import mercadopago
import pandas as pd
import requests
from flask import (
    Flask,
    jsonify,
    redirect,
    render_template,
    request,
    send_file,
    url_for,
    session,
    make_response,
)
from flask_cors import CORS
from openai import OpenAI
from PIL import Image, ImageOps, ImageDraw
import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as ReportLabImage
from reportlab.lib import colors
from reportlab.lib.units import inch
import io

# ==============================
# Configuración
# ==============================

class Config:
    # Seguridad y servidor
    SECRET_KEY = os.getenv("SECRET_KEY", "inertiax_secret_key_2025")
    MAX_CONTENT_LENGTH = int(os.getenv("MAX_CONTENT_LENGTH_MB", "20")) * 1024 * 1024
    SESSION_COOKIE_NAME = "inertiax_session"

    # Archivos
    UPLOAD_DIR = os.getenv("UPLOAD_DIR", "/tmp")
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    ALLOWED_EXT = {".csv", ".xls", ".xlsx"}

    # Pago / IA / Dominio
    DOMAIN_URL = os.getenv("DOMAIN_URL", "https://inertiax-calculadora-1.onrender.com")
    MP_ACCESS_TOKEN = os.getenv("MP_ACCESS_TOKEN")
    MP_PUBLIC_KEY = os.getenv("MP_PUBLIC_KEY")
    
    # OpenAI Config
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
    OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o")

    # Códigos de invitado
    GUEST_CODES = set(
        (os.getenv("GUEST_CODES") or "INERTIAXVIP2025,ENTRENADORPRO,INVEXORTEST").split(",")
    )

# ==============================
# App y extensiones
# ==============================

app = Flask(__name__, static_folder="static", template_folder="templates")
app.config.from_object(Config)
app.secret_key = app.config["SECRET_KEY"]

# CORS para permitir uso embebido (Shopify)
CORS(app, resources={r"/*": {"origins": "*"}}, supports_credentials=True)

# Logging
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s",
)
log = logging.getLogger("inertiax")

# Clientes externos
mp = mercadopago.SDK(app.config["MP_ACCESS_TOKEN"]) if app.config["MP_ACCESS_TOKEN"] else None
ai_client = OpenAI(api_key=app.config["OPENAI_API_KEY"]) if app.config["OPENAI_API_KEY"] else None

# ==============================
# Modelos de estado (en disco)
# ==============================

@dataclass
class JobState:
    job_id: str
    file_path: Optional[str] = None
    file_name: Optional[str] = None
    payment_ok: bool = False
    meta_path: Optional[str] = None

def _job_dir(job_id: str) -> str:
    d = os.path.join(app.config["UPLOAD_DIR"], job_id)
    os.makedirs(d, exist_ok=True)
    return d

def _job_meta_path(job_id: str) -> str:
    return os.path.join(_job_dir(job_id), "meta.json")

def _save_meta(job_id: str, meta: Dict) -> str:
    p = _job_meta_path(job_id)
    with open(p, "w", encoding="utf-8") as f:
        json.dump(meta, f, ensure_ascii=False)
    return p

def _load_meta(job_id: str) -> Dict:
    p = _job_meta_path(job_id)
    if not os.path.exists(p):
        return {}
    with open(p, "r", encoding="utf-8") as f:
        return json.load(f)

def _ensure_job() -> str:
    """Obtiene/crea un job_id mínimo en cookie de sesión."""
    if "job_id" not in session:
        session["job_id"] = uuid.uuid4().hex
    return session["job_id"]

def _allowed_file(name: str) -> bool:
    ext = os.path.splitext(name)[1].lower()
    return ext in app.config["ALLOWED_EXT"]

# ==============================
# Utilidades básicas
# ==============================

def parse_dataframe(path: str) -> pd.DataFrame:
    ext = os.path.splitext(path)[1].lower()
    if ext == ".csv":
        return pd.read_csv(path)
    return pd.read_excel(path)

def preprocess_data_by_origin(df: pd.DataFrame, origin: str) -> pd.DataFrame:
    """
    Ajusta el DataFrame según la app o dispositivo que generó el CSV.
    """
    origin = origin.lower()
    if origin == "app_android_encoder_vertical":
        rename_map = {
            "Athlete": "atleta",
            "Exercise": "ejercicio",
            "Date": "fecha",
            "Repetition": "repeticion",
            "Load(kg)": "carga_kg",
            "ConcentricVelocity(m/s)": "velocidad_concentrica_m_s",
            "EccentricVelocity(m/s)": "velocidad_eccentrica_m_s",
            "MaxVelocity(m/s)": "velocidad_maxima_m_s",
            "Duration(s)": "duracion_s",
            "Estimated1RM": "estimado_1rm_kg"
        }
        df.rename(columns={c: rename_map.get(c, c) for c in df.columns}, inplace=True)

        num_cols = [
            "carga_kg", "velocidad_concentrica_m_s", "velocidad_eccentrica_m_s",
            "velocidad_maxima_m_s", "duracion_s", "estimado_1rm_kg"
        ]
        for col in num_cols:
            if col in df.columns:
                df[col] = pd.to_numeric(df[col], errors="coerce")

        if "velocidad_concentrica_m_s" in df.columns and "carga_kg" in df.columns:
            df["potencia_relativa_w_kg"] = df["carga_kg"] * df["velocidad_concentrica_m_s"]

        if "fecha" in df.columns:
            df["fecha"] = pd.to_datetime(df["fecha"], errors="coerce")

        df.dropna(how="all", inplace=True)
        return df

    elif origin == "jump_sensor":
        if "altura_cm" in df.columns:
            df["altura_m"] = df["altura_cm"] / 100
        return df

    else:
        df.columns = [c.strip().lower().replace(" ", "_") for c in df.columns]
        return df

# ==============================
# ANÁLISIS ESPECÍFICO PARA ENCODER VERTICAL
# ==============================

def analyze_encoder_vertical_data(df: pd.DataFrame) -> Dict:
    """
    Análisis especializado para datos del encoder vertical Android
    """
    analysis = {
        "atletas": {},
        "resumen_general": {},
        "graficos": []
    }
    
    # Identificar atletas únicos
    if "atleta" not in df.columns:
        # Si no hay columna atleta, asumir un solo atleta
        atletas = ["Atleta Principal"]
        df["atleta"] = "Atleta Principal"
    else:
        atletas = df["atleta"].unique()
    
    analysis["resumen_general"]["total_atletas"] = len(atletas)
    analysis["resumen_general"]["atletas_identificados"] = list(atletas)
    
    # Análisis por atleta
    for atleta in atletas:
        df_atleta = df[df["atleta"] == atleta].copy()
        
        atleta_analysis = {
            "ejercicios": {},
            "metricas_generales": {},
            "datos_brutos": df_atleta.to_dict('records')
        }
        
        # Identificar ejercicios únicos
        ejercicios = df_atleta["ejercicio"].unique() if "ejercicio" in df_atleta.columns else ["Ejercicio Principal"]
        atleta_analysis["metricas_generales"]["total_ejercicios"] = len(ejercicios)
        atleta_analysis["metricas_generales"]["ejercicios_identificados"] = list(ejercicios)
        
        # Análisis por ejercicio
        for ejercicio in ejercicios:
            df_ejercicio = df_atleta[df_atleta["ejercicio"] == ejercicio] if "ejercicio" in df_atleta.columns else df_atleta
            
            ejercicio_analysis = {
                "repeticiones": {},
                "series": {},
                "metricas": {}
            }
            
            # Calcular métricas básicas
            if "carga_kg" in df_ejercicio.columns:
                ejercicio_analysis["metricas"]["carga_promedio"] = df_ejercicio["carga_kg"].mean()
                ejercicio_analysis["metricas"]["carga_maxima"] = df_ejercicio["carga_kg"].max()
                ejercicio_analysis["metricas"]["carga_minima"] = df_ejercicio["carga_kg"].min()
            
            if "velocidad_concentrica_m_s" in df_ejercicio.columns:
                ejercicio_analysis["metricas"]["vel_concentrica_promedio"] = df_ejercicio["velocidad_concentrica_m_s"].mean()
                ejercicio_analysis["metricas"]["vel_concentrica_maxima"] = df_ejercicio["velocidad_concentrica_m_s"].max()
            
            if "velocidad_eccentrica_m_s" in df_ejercicio.columns:
                ejercicio_analysis["metricas"]["vel_excentrica_promedio"] = df_ejercicio["velocidad_eccentrica_m_s"].mean()
            
            if "repeticion" in df_ejercicio.columns:
                ejercicio_analysis["metricas"]["total_repeticiones"] = df_ejercicio["repeticion"].max()
            
            if "estimado_1rm_kg" in df_ejercicio.columns:
                ejercicio_analysis["metricas"]["1rm_estimado_promedio"] = df_ejercicio["estimado_1rm_kg"].mean()
                ejercicio_analysis["metricas"]["1rm_estimado_maximo"] = df_ejercicio["estimado_1rm_kg"].max()
            
            atleta_analysis["ejercicios"][ejercicio] = ejercicio_analysis
        
        # Métricas generales del atleta
        if "carga_kg" in df_atleta.columns:
            atleta_analysis["metricas_generales"]["carga_total_promedio"] = df_atleta["carga_kg"].mean()
        
        if "velocidad_concentrica_m_s" in df_atleta.columns:
            atleta_analysis["metricas_generales"]["velocidad_promedio"] = df_atleta["velocidad_concentrica_m_s"].mean()
        
        analysis["atletas"][atleta] = atleta_analysis
    
    return analysis

def generate_encoder_vertical_charts(df: pd.DataFrame, analysis: Dict) -> List[BytesIO]:
    """
    Genera gráficos específicos para datos de encoder vertical
    """
    charts = []
    
    try:
        # Configuración general de matplotlib
        plt.style.use('seaborn-v0_8')
        
        # Gráfico 1: Velocidad vs Carga por atleta
        if "atleta" in df.columns and "carga_kg" in df.columns and "velocidad_concentrica_m_s" in df.columns:
            fig, ax = plt.subplots(figsize=(10, 6))
            
            for atleta in df["atleta"].unique():
                df_atleta = df[df["atleta"] == atleta]
                ax.scatter(df_atleta["carga_kg"], df_atleta["velocidad_concentrica_m_s"], 
                          label=atleta, alpha=0.7, s=50)
            
            ax.set_xlabel('Carga (kg)')
            ax.set_ylabel('Velocidad Concéntrica (m/s)')
            ax.set_title('Relación Velocidad vs Carga por Atleta')
            ax.legend()
            ax.grid(True, alpha=0.3)
            
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            charts.append(buf)
            plt.close()
        
        # Gráfico 2: Perfil Fuerza-Velocidad por ejercicio
        if "ejercicio" in df.columns and "carga_kg" in df.columns and "velocidad_concentrica_m_s" in df.columns:
            ejercicios = df["ejercicio"].unique()[:4]  # Máximo 4 ejercicios
            fig, axes = plt.subplots(2, 2, figsize=(12, 10))
            axes = axes.ravel()
            
            for i, ejercicio in enumerate(ejercicios):
                if i >= 4:
                    break
                    
                df_ejercicio = df[df["ejercicio"] == ejercicio]
                axes[i].scatter(df_ejercicio["carga_kg"], df_ejercicio["velocidad_concentrica_m_s"], 
                               alpha=0.7, color='blue')
                
                # Ajustar línea de tendencia
                if len(df_ejercicio) > 1:
                    z = np.polyfit(df_ejercicio["carga_kg"], df_ejercicio["velocidad_concentrica_m_s"], 1)
                    p = np.poly1d(z)
                    axes[i].plot(df_ejercicio["carga_kg"], p(df_ejercicio["carga_kg"]), 
                                "r--", alpha=0.8, label=f"Tendencia")
                
                axes[i].set_xlabel('Carga (kg)')
                axes[i].set_ylabel('Velocidad (m/s)')
                axes[i].set_title(f'Perfil F-V: {ejercicio}')
                axes[i].legend()
                axes[i].grid(True, alpha=0.3)
            
            # Ocultar ejes vacíos si hay menos de 4 ejercicios
            for i in range(len(ejercicios), 4):
                axes[i].set_visible(False)
            
            plt.tight_layout()
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            charts.append(buf)
            plt.close()
        
        # Gráfico 3: Evolución de velocidad por repetición
        if "repeticion" in df.columns and "velocidad_concentrica_m_s" in df.columns:
            fig, ax = plt.subplots(figsize=(10, 6))
            
            if "atleta" in df.columns:
                for atleta in df["atleta"].unique()[:3]:  # Máximo 3 atletas
                    df_atleta = df[df["atleta"] == atleta]
                    velocidad_por_rep = df_atleta.groupby("repeticion")["velocidad_concentrica_m_s"].mean()
                    ax.plot(velocidad_por_rep.index, velocidad_por_rep.values, 
                           marker='o', label=atleta, linewidth=2)
            else:
                velocidad_por_rep = df.groupby("repeticion")["velocidad_concentrica_m_s"].mean()
                ax.plot(velocidad_por_rep.index, velocidad_por_rep.values, 
                       marker='o', color='blue', linewidth=2, label='Velocidad Promedio')
            
            ax.set_xlabel('Número de Repetición')
            ax.set_ylabel('Velocidad Concéntrica (m/s)')
            ax.set_title('Evolución de Velocidad por Repetición')
            ax.legend()
            ax.grid(True, alpha=0.3)
            
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            charts.append(buf)
            plt.close()
        
        # Gráfico 4: Comparación de 1RM estimado entre atletas
        if "atleta" in df.columns and "estimado_1rm_kg" in df.columns:
            fig, ax = plt.subplots(figsize=(10, 6))
            
            datos_1rm = df.groupby("atleta")["estimado_1rm_kg"].mean().sort_values(ascending=False)
            bars = ax.bar(datos_1rm.index, datos_1rm.values, color='skyblue', alpha=0.7)
            
            # Añadir valores en las barras
            for bar in bars:
                height = bar.get_height()
                ax.text(bar.get_x() + bar.get_width()/2., height,
                       f'{height:.1f} kg',
                       ha='center', va='bottom')
            
            ax.set_xlabel('Atleta')
            ax.set_ylabel('1RM Estimado (kg)')
            ax.set_title('Comparación de 1RM Estimado entre Atletas')
            plt.xticks(rotation=45)
            ax.grid(True, alpha=0.3, axis='y')
            
            buf = BytesIO()
            plt.savefig(buf, format='png', dpi=150, bbox_inches='tight')
            buf.seek(0)
            charts.append(buf)
            plt.close()
            
    except Exception as e:
        log.error(f"Error generando gráficos encoder vertical: {e}")
    
    return charts

def generate_encoder_vertical_pdf(analysis: Dict, charts: List[BytesIO], meta: Dict) -> str:
    """
    Genera PDF especializado para datos de encoder vertical
    """
    try:
        pdf_path = os.path.join(app.config["UPLOAD_DIR"], f"reporte_encoder_vertical_{uuid.uuid4().hex}.pdf")
        
        doc = SimpleDocTemplate(pdf_path, pagesize=A4, topMargin=1*inch)
        styles = getSampleStyleSheet()
        story = []
        
        # Título principal
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1,  # Centrado
            textColor=colors.HexColor('#2E86AB')
        )
        
        story.append(Paragraph("REPORTE ENCODER VERTICAL - ANÁLISIS COMPLETO", title_style))
        story.append(Spacer(1, 20))
        
        # Información general
        story.append(Paragraph(f"Entrenador: {meta.get('nombre_entrenador', 'No especificado')}", styles['Heading2']))
        story.append(Paragraph(f"Fecha de generación: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Resumen general
        story.append(Paragraph("RESUMEN GENERAL", styles['Heading2']))
        resumen = analysis["resumen_general"]
        story.append(Paragraph(f"Total de atletas analizados: {resumen['total_atletas']}", styles['Normal']))
        story.append(Paragraph(f"Atletas identificados: {', '.join(resumen['atletas_identificados'])}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Análisis individual por atleta
        for atleta, datos_atleta in analysis["atletas"].items():
            story.append(Paragraph(f"ANÁLISIS INDIVIDUAL: {atleta}", styles['Heading2']))
            story.append(Spacer(1, 10))
            
            # Métricas generales del atleta
            metricas = datos_atleta["metricas_generales"]
            story.append(Paragraph(f"Ejercicios realizados: {metricas['total_ejercicios']}", styles['Normal']))
            story.append(Paragraph(f"Ejercicios: {', '.join(metricas['ejercicios_identificados'])}", styles['Normal']))
            
            if "carga_total_promedio" in metricas:
                story.append(Paragraph(f"Carga promedio: {metricas['carga_total_promedio']:.1f} kg", styles['Normal']))
            
            if "velocidad_promedio" in metricas:
                story.append(Paragraph(f"Velocidad concéntrica promedio: {metricas['velocidad_promedio']:.3f} m/s", styles['Normal']))
            
            story.append(Spacer(1, 15))
            
            # Análisis por ejercicio
            for ejercicio, datos_ejercicio in datos_atleta["ejercicios"].items():
                story.append(Paragraph(f"Ejercicio: {ejercicio}", styles['Heading3']))
                
                metricas_ej = datos_ejercicio["metricas"]
                datos_tabla = []
                
                if "carga_promedio" in metricas_ej:
                    datos_tabla.append(["Carga promedio", f"{metricas_ej['carga_promedio']:.1f} kg"])
                if "carga_maxima" in metricas_ej:
                    datos_tabla.append(["Carga máxima", f"{metricas_ej['carga_maxima']:.1f} kg"])
                if "vel_concentrica_promedio" in metricas_ej:
                    datos_tabla.append(["Velocidad concéntrica promedio", f"{metricas_ej['vel_concentrica_promedio']:.3f} m/s"])
                if "vel_concentrica_maxima" in metricas_ej:
                    datos_tabla.append(["Velocidad concéntrica máxima", f"{metricas_ej['vel_concentrica_maxima']:.3f} m/s"])
                if "1rm_estimado_promedio" in metricas_ej:
                    datos_tabla.append(["1RM estimado promedio", f"{metricas_ej['1rm_estimado_promedio']:.1f} kg"])
                if "total_repeticiones" in metricas_ej:
                    datos_tabla.append(["Total de repeticiones", f"{metricas_ej['total_repeticiones']}"])
                
                if datos_tabla:
                    tabla = Table(datos_tabla, colWidths=[3*inch, 2*inch])
                    tabla.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#F0F0F0')),
                        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ]))
                    story.append(tabla)
                
                story.append(Spacer(1, 10))
            
            story.append(Spacer(1, 20))
        
        # Añadir gráficos al PDF
        if charts:
            story.append(Paragraph("GRÁFICOS DE ANÁLISIS", styles['Heading2']))
            story.append(Spacer(1, 20))
            
            for i, chart in enumerate(charts[:4]):  # Máximo 4 gráficos
                try:
                    chart.seek(0)
                    img = ReportLabImage(chart, width=6*inch, height=4*inch)
                    story.append(img)
                    story.append(Spacer(1, 10))
                    story.append(Paragraph(f"Gráfico {i+1}", styles['Italic']))
                    story.append(Spacer(1, 20))
                except Exception as e:
                    log.error(f"Error añadiendo gráfico {i} al PDF: {e}")
                    continue
        
        doc.build(story)
        return pdf_path
        
    except Exception as e:
        log.error(f"Error generando PDF encoder vertical: {e}")
        # PDF de error simplificado
        error_path = os.path.join(app.config["UPLOAD_DIR"], f"error_{uuid.uuid4().hex}.pdf")
        c = canvas.Canvas(error_path)
        c.drawString(100, 750, "Error generando reporte encoder vertical")
        c.drawString(100, 730, str(e))
        c.save()
        return error_path

# ==============================
# NÚCLEO IA (modificado para encoder vertical)
# ==============================

def run_complete_ai_analysis(df: pd.DataFrame, meta: dict) -> dict:
    """
    Análisis completo por IA con enfoque en encoder vertical
    """
    if not ai_client:
        return {
            "analysis": "⚠️ Servicio de IA no disponible. Configure OPENAI_API_KEY.",
            "python_code_for_charts": "",
            "charts_description": "",
            "recommendations": ["Configure la API key de OpenAI para habilitar el análisis IA"]
        }
    
    # Primero realizar análisis específico de encoder vertical
    encoder_analysis = analyze_encoder_vertical_data(df)
    
    n_rows, n_cols = df.shape
    data_preview = df.head(500).to_csv(index=False)
    
    contexto = "\n".join([
        "ANÁLISIS ESPECIALIZADO EN ENCODER VERTICAL - DATOS DE FUERZA-VELOCIDAD",
        f"Entrenador: {meta.get('nombre_entrenador', 'No especificado')}",
        f"Origen: Encoder Vertical Android",
        f"Total atletas: {encoder_analysis['resumen_general']['total_atletas']}",
        f"Atletas: {', '.join(encoder_analysis['resumen_general']['atletas_identificados'])}",
        f"Dataset: {n_rows} filas × {n_cols} columnas",
        f"Columnas disponibles: {', '.join(df.columns.tolist())}",
        "",
        "INSTRUCCIONES ESPECÍFICAS:",
        "1. Enfócate en el análisis de perfiles fuerza-velocidad",
        "2. Analiza la relación carga-velocidad para cada atleta",
        "3. Identifica fatiga mediante evolución de velocidad por repetición",
        "4. Evalúa calidad técnica mediante consistencia de velocidades",
        "5. Proporciona recomendaciones específicas por atleta",
        "6. Considera métricas como: velocidad media, 1RM estimado, potencia"
    ])

    system_prompt = """
Eres un especialista en biomecánica y análisis de datos de fuerza-velocidad usando encoder lineal.
Analiza los datos de encoder vertical y genera un reporte profesional enfocado en:
- Perfiles individuales fuerza-velocidad por atleta
- Evolución de la fatiga durante las series
- Calidad técnica y consistencia de movimientos
- Estimaciones de 1RM y potencia
- Recomendaciones personalizadas por atleta

Responde EXACTAMENTE en formato JSON con estas claves:
- analysis: texto completo del análisis profesional
- python_code_for_charts: código Python para generar gráficos adicionales (opcional)
- charts_description: descripción de los gráficos generados
- recommendations: lista de recomendaciones específicas por atleta
"""

    user_prompt = f"{contexto}\n\nDatos CSV (primeras filas):\n```csv\n{data_preview}\n```"

    try:
        response = ai_client.chat.completions.create(
            model=app.config["OPENAI_MODEL"],
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.3,
            max_tokens=4000,
            response_format={"type": "json_object"}
        )
        return json.loads(response.choices[0].message.content)
    except Exception as e:
        log.error(f"Error en IA: {e}")
        return {
            "analysis": f"Error en el análisis IA: {str(e)}",
            "python_code_for_charts": "",
            "charts_description": "",
            "recommendations": ["Error técnico - contacte al administrador"]
        }

def execute_ai_charts_code(python_code: str, df: pd.DataFrame) -> List[BytesIO]:
    """Ejecuta código Python de gráficos generado por IA"""
    if not python_code.strip():
        return []
    
    try:
        import matplotlib.pyplot as plt
        import seaborn as sns
        import numpy as np
        
        # Resetear cualquier figura previa
        plt.close('all')
        
        local_vars = {
            'df': df, 
            'BytesIO': BytesIO, 
            'plt': plt,
            'sns': sns,
            'np': np,
            'pd': pd,
            'charts': []
        }
        
        # Ejecutar el código
        exec(python_code, local_vars)
        
        # Obtener gráficos
        charts = local_vars.get('charts', [])
        if not isinstance(charts, list):
            charts = []
            
        return charts
        
    except Exception as e:
        log.error(f"Error ejecutando código de gráficos: {e}")
        return []

def generate_pdf_from_ai(ai_result: dict, charts: List[BytesIO], meta: dict) -> str:
    """Genera PDF con el análisis de IA"""
    try:
        pdf_path = os.path.join(app.config["UPLOAD_DIR"], f"reporte_ia_{uuid.uuid4().hex}.pdf")
        c = canvas.Canvas(pdf_path, pagesize=letter)
        
        # Header
        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, 750, "Reporte InertiaX - Análisis IA")
        c.setFont("Helvetica", 10)
        c.drawString(100, 730, f"Entrenador: {meta.get('nombre_entrenador', '-')}")
        c.drawString(100, 715, f"Origen: Encoder Vertical Android")
        c.drawString(100, 700, f"Fecha: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
        
        # Análisis
        c.setFont("Helvetica-Bold", 12)
        c.drawString(50, 680, "ANÁLISIS PROFESIONAL:")
        c.setFont("Helvetica", 9)
        
        # Text wrapping simple
        text = ai_result.get('analysis', 'No analysis available')
        y = 660
        lines = []
        words = text.split()
        line = ""
        
        for word in words:
            test_line = f"{line} {word}".strip()
            if len(test_line) < 80:
                line = test_line
            else:
                lines.append(line)
                line = word
        if line:
            lines.append(line)
        
        for line in lines[:30]:  # Máximo 30 líneas
            if y < 100:
                c.showPage()
                y = 750
            c.drawString(50, y, line)
            y -= 12
        
        # Gráficos
        if charts:
            c.showPage()
            c.setFont("Helvetica-Bold", 14)
            c.drawString(50, 750, "GRÁFICOS DE ANÁLISIS")
            
            y_pos = 700
            for i, chart in enumerate(charts[:3]):  # Máximo 3 gráficos
                try:
                    chart.seek(0)
                    c.drawImage(ImageReader(chart), 50, y_pos - 250, width=500, height=200)
                    y_pos -= 280
                    if y_pos < 100 and i < len(charts) - 1:
                        c.showPage()
                        y_pos = 750
                except Exception as e:
                    log.error(f"Error dibujando gráfico {i}: {e}")
        
        # Recomendaciones
        c.showPage()
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, 750, "RECOMENDACIONES")
        c.setFont("Helvetica", 10)
        
        recommendations = ai_result.get('recommendations', [])
        y_rec = 720
        if isinstance(recommendations, list):
            for rec in recommendations[:10]:  # Máximo 10 recomendaciones
                if y_rec < 50:
                    c.showPage()
                    y_rec = 750
                c.drawString(70, y_rec, f"• {rec}")
                y_rec -= 20
        else:
            c.drawString(50, 720, str(recommendations))
        
        c.save()
        return pdf_path
        
    except Exception as e:
        log.error(f"Error generando PDF: {e}")
        # Crear PDF de error
        error_path = os.path.join(app.config["UPLOAD_DIR"], f"error_{uuid.uuid4().hex}.pdf")
        c = canvas.Canvas(error_path)
        c.drawString(100, 750, "Error generando reporte")
        c.drawString(100, 730, str(e))
        c.save()
        return error_path

def generate_docx_from_ai(ai_result: dict, charts: List[BytesIO], meta: dict) -> str:
    """Genera DOCX con el análisis de IA"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        
        # Título
        doc.add_heading('Reporte InertiaX - Análisis Encoder Vertical', 0)
        doc.add_paragraph(f"Entrenador: {meta.get('nombre_entrenador', '-')}")
        doc.add_paragraph(f"Origen: Encoder Vertical Android")
        doc.add_paragraph(f"Fecha: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
        
        # Análisis
        doc.add_heading('Análisis Profesional', level=1)
        doc.add_paragraph(ai_result.get('analysis', 'No analysis available'))
        
        # Gráficos
        if charts:
            doc.add_heading('Gráficos', level=1)
            for i, chart in enumerate(charts[:3]):
                try:
                    chart.seek(0)
                    doc.add_picture(chart, width=Inches(6.0))
                    doc.add_paragraph(f"Gráfico {i+1}")
                except Exception as e:
                    doc.add_paragraph(f"Error cargando gráfico {i+1}: {e}")
        
        # Recomendaciones
        doc.add_heading('Recomendaciones', level=1)
        recommendations = ai_result.get('recommendations', [])
        if isinstance(recommendations, list):
            for rec in recommendations:
                doc.add_paragraph(f"• {rec}", style='List Bullet')
        else:
            doc.add_paragraph(str(recommendations))
        
        path = os.path.join(app.config["UPLOAD_DIR"], f"reporte_ia_{uuid.uuid4().hex}.docx")
        doc.save(path)
        return path
        
    except Exception as e:
        log.error(f"Error generando DOCX: {e}")
        # Crear DOCX simple de error
        doc = Document()
        doc.add_heading('Error', 0)
        doc.add_paragraph(f"Error generando reporte: {e}")
        error_path = os.path.join(app.config["UPLOAD_DIR"], f"error_{uuid.uuid4().hex}.docx")
        doc.save(error_path)
        return error_path

# ==============================
# RUTAS PRINCIPALES (SIMPLIFICADAS)
# ==============================

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/health")
def health():
    return jsonify({"status": "ok", "message": "InertiaX API is running"})



@app.route("/upload", methods=["POST"])
def upload():
    """Subida de archivo y procesamiento inicial - INTERFAZ SIMPLIFICADA"""
    try:
        job_id = _ensure_job()
        session.modified = True

        # Formulario simplificado
        form = {
            "nombre_entrenador": request.form.get("nombre_entrenador", "").strip(),
            "origen_app": request.form.get("origen_app", "").strip(),
            "codigo_invitado": request.form.get("codigo_invitado", "").strip(),
        }

        # Verificar código de invitado
        code = form.get("codigo_invitado", "").strip()
        payment_ok = False
        mensaje = None
        if code and code in app.config["GUEST_CODES"]:
            payment_ok = True
            mensaje = "🔓 Código de invitado válido. Puedes generar tu reporte."

        f = request.files.get("file")
        if not f or f.filename == "":
            return render_template("index.html", error="No se subió ningún archivo.")

        if not _allowed_file(f.filename):
            return render_template("index.html", error="Formato no permitido. Usa .csv, .xls o .xlsx")

        ext = os.path.splitext(f.filename)[1].lower()
        safe_name = f"{uuid.uuid4().hex}{ext}"
        save_path = os.path.join(_job_dir(job_id), safe_name)
        f.save(save_path)

        # Persistir meta
        meta = {
            "file_name": f.filename,
            "file_path": save_path,
            "payment_ok": payment_ok,
            "form": form,
        }
        _save_meta(job_id, meta)

        # Previsualización simple
        try:
            df = parse_dataframe(save_path)
            df = preprocess_data_by_origin(df, form.get("origen_app", ""))
            
            # Generar tabla HTML para previsualización
            table_html = df.head(10).to_html(
                classes="table table-striped table-bordered table-hover", 
                index=False,
                escape=False
            )
            
            return render_template(
                "index.html",
                table_html=table_html,
                filename=f.filename,
                form_data=form,
                mensaje=mensaje,
                show_payment=(not payment_ok),
            )
        except Exception as e:
            log.exception("Error al procesar archivo")
            return render_template("index.html", error=f"Error al procesar el archivo: {str(e)}")
            
    except Exception as e:
        log.exception("Error general en upload")
        return render_template("index.html", error=f"Error en el servidor: {str(e)}")




@app.route("/create_preference", methods=["POST"])
def create_preference():
    """Crea preferencia Mercado Pago"""
    if not mp:
        return jsonify(error="Mercado Pago no configurado"), 500

    job_id = session.get("job_id")
    if not job_id:
        return jsonify(error="Sesión inválida"), 400

    meta = _load_meta(job_id)
    form = meta.get("form", {})

    # Precio fijo simplificado
    price = 4900  # Precio único para análisis encoder vertical

    pref_data = {
        "items": [{
            "title": f"InertiaX - Análisis Encoder Vertical",
            "quantity": 1,
            "unit_price": price,
            "currency_id": "CLP",
        }],
        "back_urls": {
            "success": f"{app.config['DOMAIN_URL']}/success",
            "failure": f"{app.config['DOMAIN_URL']}/cancel", 
            "pending": f"{app.config['DOMAIN_URL']}/cancel",
        },
        "auto_return": "approved",
    }

    try:
        pref = mp.preference().create(pref_data)
        return jsonify(pref.get("response", {}))
    except Exception as e:
        log.exception("Error creando preferencia MP")
        return jsonify(error=str(e)), 500

@app.route("/success")
def success():
    """Pago exitoso"""
    job_id = session.get("job_id")
    if not job_id:
        return redirect(url_for("index"))

    meta = _load_meta(job_id)
    meta["payment_ok"] = True
    _save_meta(job_id, meta)
    return render_template("success.html")

@app.route("/cancel") 
def cancel():
    """Pago cancelado"""
    return render_template("cancel.html")

@app.route("/generate_report")
def generate_report():
    """Genera y descarga el reporte completo con análisis encoder vertical"""
    job_id = session.get("job_id")
    if not job_id:
        return redirect(url_for("index"))

    meta = _load_meta(job_id)
    if not meta.get("payment_ok"):
        return redirect(url_for("index"))

    file_path = meta.get("file_path")
    if not file_path or not os.path.exists(file_path):
        return render_template("index.html", error="No hay archivo para analizar.")

    try:
        # 1. Cargar datos
        df = parse_dataframe(file_path)
        df = preprocess_data_by_origin(df, meta.get("form", {}).get("origen_app", ""))
        
        # 2. Análisis específico encoder vertical
        log.info("Realizando análisis encoder vertical...")
        encoder_analysis = analyze_encoder_vertical_data(df)
        
        # 3. Generar gráficos específicos encoder vertical
        log.info("Generando gráficos encoder vertical...")
        encoder_charts = generate_encoder_vertical_charts(df, encoder_analysis)
        
        # 4. Generar PDF especializado encoder vertical
        pdf_path = generate_encoder_vertical_pdf(encoder_analysis, encoder_charts, meta.get("form", {}))
        
        # 5. Análisis IA adicional si está disponible
        ai_result = {}
        ai_charts = []
        if ai_client:
            log.info("Ejecutando análisis IA complementario...")
            ai_result = run_complete_ai_analysis(df, meta.get("form", {}))
            
            python_code = ai_result.get("python_code_for_charts", "")
            if python_code:
                ai_charts = execute_ai_charts_code(python_code, df)
            
            # Generar PDF de IA
            ai_pdf_path = generate_pdf_from_ai(ai_result, ai_charts, meta.get("form", {}))
            ai_docx_path = generate_docx_from_ai(ai_result, ai_charts, meta.get("form", {}))
        else:
            ai_pdf_path = None
            ai_docx_path = None
        
        # 6. Crear ZIP con todos los reportes
        zip_path = os.path.join(_job_dir(job_id), f"reporte_completo_{uuid.uuid4().hex}.zip")
        with zipfile.ZipFile(zip_path, "w") as zf:
            # Reporte principal encoder vertical
            zf.write(pdf_path, "reporte_encoder_vertical.pdf")
            
            # Reportes IA si están disponibles
            if ai_pdf_path:
                zf.write(ai_pdf_path, "analisis_ia_complementario.pdf")
            if ai_docx_path:
                zf.write(ai_docx_path, "analisis_ia_complementario.docx")
            
            # Datos originales
            zf.write(file_path, os.path.basename(meta.get("file_name", "datos_original.csv")))
        
        # Limpiar archivos temporales
        try:
            os.remove(pdf_path)
            if ai_pdf_path:
                os.remove(ai_pdf_path)
            if ai_docx_path:
                os.remove(ai_docx_path)
        except:
            pass
            
        return send_file(
            zip_path, 
            as_attachment=True, 
            download_name=f"reporte_encoder_vertical_{uuid.uuid4().hex[:8]}.zip"
        )
        
    except Exception as e:
        log.exception("Error generando reporte")
        return render_template("index.html", error=f"Error generando reporte: {e}")

@app.route("/preview_analysis")
def preview_analysis():
    """Vista previa del análisis encoder vertical"""
    job_id = session.get("job_id")
    if not job_id:
        return redirect(url_for("index"))

    meta = _load_meta(job_id)
    if not meta.get("payment_ok"):
        return redirect(url_for("index"))

    file_path = meta.get("file_path")
    if not file_path or not os.path.exists(file_path):
        return render_template("index.html", error="No hay archivo para analizar.")

    try:
        df = parse_dataframe(file_path)
        df = preprocess_data_by_origin(df, meta.get("form", {}).get("origen_app", ""))
        
        # Análisis encoder vertical
        encoder_analysis = analyze_encoder_vertical_data(df)
        
        # Análisis IA si está disponible
        ai_result = {}
        if ai_client:
            ai_result = run_complete_ai_analysis(df, meta.get("form", {}))
        
        return render_template(
            "preview.html",
            encoder_analysis=encoder_analysis,
            ai_analysis=ai_result.get("analysis", "Análisis IA no disponible"),
            recommendations=ai_result.get("recommendations", []),
            filename=meta.get("file_name")
        )
        
    except Exception as e:
        log.exception("Error en vista previa")
        return render_template("index.html", error=f"Error en vista previa: {e}")

# ==============================
# Manejo de errores
# ==============================

@app.errorhandler(413)
def too_large(_e):
    return render_template("index.html", error="Archivo demasiado grande.")

@app.errorhandler(404)
def not_found(_e):
    return render_template("index.html", error="Página no encontrada.")

@app.errorhandler(500)
def internal_error(_e):
    return render_template("index.html", error="Error interno del servidor.")

@app.errorhandler(Exception)
def global_error(e):
    log.exception("Error no controlado")
    return render_template("index.html", error=f"Error interno: {e}")

# ==============================
# Run
# ==============================

if __name__ == "__main__":
    port = int(os.environ.get("PORT", "10000"))
    app.run(host="0.0.0.0", port=port)
